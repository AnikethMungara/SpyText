"""
DOCX text extraction with metadata.

Extracts text from Microsoft Word documents (.docx) with positioning and formatting metadata.
"""

from pathlib import Path
from typing import List, Optional, Tuple
import logging

from docx import Document
from docx.shared import RGBColor
from docx.text.run import Run
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from src.models import TextSpan

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """
    Extract text and metadata from DOCX files.

    Responsibilities:
    - Extract text with formatting metadata
    - Capture font size, color, and styling
    - Preserve document structure (paragraphs, tables, text boxes)
    - Handle page breaks for accurate page numbering
    - Extract text from shapes and text boxes
    """

    # Default background color for DOCX (white)
    DEFAULT_BACKGROUND = (255, 255, 255)

    # Estimated lines per page for page calculation
    LINES_PER_PAGE = 40

    def __init__(self, config: Optional[dict] = None):
        """
        Initialize DOCX extractor.

        Args:
            config: Configuration dict from settings.yaml
        """
        self.config = config or {}
        self.current_page = 1
        self.line_count = 0

    def extract(self, docx_path: Path) -> List[TextSpan]:
        """
        Extract all text spans from a DOCX file with metadata.

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of TextSpan objects with formatting metadata

        Raises:
            ValueError: If DOCX is invalid or corrupted
        """
        logger.info(f"Extracting text from DOCX: {docx_path}")

        try:
            doc = Document(str(docx_path))
            spans = []

            # Reset page tracking
            self.current_page = 1
            self.line_count = 0

            # Process document body elements (paragraphs, tables, etc.)
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Regular paragraph
                    paragraph = Paragraph(element, doc)
                    para_spans = self._extract_from_paragraph(paragraph)
                    spans.extend(para_spans)

                    # Track page breaks
                    self._check_for_page_break(paragraph)

                elif isinstance(element, CT_Tbl):
                    # Table
                    table = Table(element, doc)
                    table_spans = self._extract_from_table(table)
                    spans.extend(table_spans)

            # Extract from headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    header_spans = self._extract_from_header_footer(section.header, "header")
                    spans.extend(header_spans)

                # Footers
                if section.footer:
                    footer_spans = self._extract_from_header_footer(section.footer, "footer")
                    spans.extend(footer_spans)

            # Extract from text boxes and shapes
            shape_spans = self._extract_from_shapes(doc)
            spans.extend(shape_spans)

            logger.info(f"Extracted {len(spans)} text spans from {docx_path.name}")
            return spans

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}") from e

    def _extract_from_paragraph(self, paragraph: Paragraph) -> List[TextSpan]:
        """Extract text spans from a paragraph."""
        spans = []

        for run in paragraph.runs:
            text = run.text

            # Skip empty runs
            if not text or not text.strip():
                continue

            # Extract formatting metadata
            font_size = self._get_font_size(run)
            font_color = self._get_font_color(run)
            bg_color = self._get_background_color(run)

            # Calculate approximate position (in points)
            y_position = self.line_count * 14  # Approximate 14pt line height

            # Create TextSpan with page-aware positioning
            span = TextSpan(
                text=text,
                page_number=self.current_page,
                bbox=(0, y_position, 500, y_position + 14),  # Approximate width and height
                font_size=font_size,
                font_color=font_color,
                background_color=bg_color
            )
            spans.append(span)

        # Increment line count
        self.line_count += 1

        return spans

    def _extract_from_table(self, table: Table) -> List[TextSpan]:
        """Extract text spans from a table."""
        spans = []

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    para_spans = self._extract_from_paragraph(paragraph)
                    spans.extend(para_spans)

        return spans

    def _extract_from_header_footer(self, header_or_footer, location: str) -> List[TextSpan]:
        """Extract text from headers or footers."""
        spans = []

        try:
            for paragraph in header_or_footer.paragraphs:
                para_spans = self._extract_from_paragraph(paragraph)
                spans.extend(para_spans)
        except Exception as e:
            logger.debug(f"Could not extract from {location}: {e}")

        return spans

    def _extract_from_shapes(self, doc: Document) -> List[TextSpan]:
        """Extract text from shapes and text boxes."""
        spans = []

        try:
            # Access inline shapes (text boxes, etc.)
            for section in doc.sections:
                # Check for floating shapes in the document
                for element in doc.element.body:
                    if hasattr(element, 'xpath'):
                        # Look for text boxes and shapes
                        textboxes = element.xpath('.//w:txbxContent//w:p')
                        for tb in textboxes:
                            paragraph = Paragraph(tb, doc)
                            para_spans = self._extract_from_paragraph(paragraph)
                            spans.extend(para_spans)
        except Exception as e:
            logger.debug(f"Could not extract from shapes: {e}")

        return spans

    def _check_for_page_break(self, paragraph: Paragraph) -> None:
        """Check if paragraph contains a page break and update page count."""
        try:
            # Check for explicit page break
            if paragraph._element.xpath('.//w:br[@w:type="page"]'):
                self.current_page += 1
                self.line_count = 0
                return

            # Estimate page break based on line count
            if self.line_count >= self.LINES_PER_PAGE:
                self.current_page += 1
                self.line_count = 0

        except Exception as e:
            logger.debug(f"Could not check for page break: {e}")

    def _get_font_size(self, run: Run) -> Optional[float]:
        """
        Extract font size from a run.

        Args:
            run: Document run object

        Returns:
            Font size in points, or None if not set
        """
        try:
            if run.font.size is not None:
                # Convert EMU (English Metric Units) to points
                # 1 point = 12700 EMU
                return run.font.size.pt
        except Exception as e:
            logger.debug(f"Could not extract font size: {e}")

        return None

    def _get_font_color(self, run: Run) -> Optional[Tuple[int, int, int]]:
        """
        Extract font color from a run.

        Args:
            run: Document run object

        Returns:
            RGB color tuple (0-255), or None if not set
        """
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                return (rgb[0], rgb[1], rgb[2])
        except Exception as e:
            logger.debug(f"Could not extract font color: {e}")

        # Default to black if not specified
        return (0, 0, 0)

    def _get_background_color(self, run: Run) -> Optional[Tuple[int, int, int]]:
        """
        Extract background/highlight color from a run.

        Args:
            run: Document run object

        Returns:
            RGB color tuple (0-255), defaults to white
        """
        try:
            # Check for highlight color
            if run.font.highlight_color:
                # Highlight colors in DOCX are named colors, not RGB
                # For simplicity, we'll map common ones
                highlight_map = {
                    'yellow': (255, 255, 0),
                    'bright_green': (0, 255, 0),
                    'turquoise': (0, 255, 255),
                    'pink': (255, 192, 203),
                    'blue': (0, 0, 255),
                    'red': (255, 0, 0),
                    'dark_blue': (0, 0, 139),
                    'teal': (0, 128, 128),
                    'green': (0, 128, 0),
                    'violet': (238, 130, 238),
                    'dark_red': (139, 0, 0),
                    'dark_yellow': (204, 204, 0),
                    'gray_50': (128, 128, 128),
                    'gray_25': (192, 192, 192),
                }

                highlight_str = str(run.font.highlight_color).lower()
                if highlight_str in highlight_map:
                    return highlight_map[highlight_str]
        except Exception as e:
            logger.debug(f"Could not extract background color: {e}")

        # Default to white background
        return self.DEFAULT_BACKGROUND
